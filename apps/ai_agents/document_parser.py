"""
文档解析器
支持解析 PDF、Word 等文档,提取文本内容用于 AI 测试用例生成
"""
import os
from PyPDF2 import PdfReader
from docx import Document

class DocumentParser:
    """文档解析器"""
    
    @staticmethod
    def parse_pdf(file_path):
        """
        解析 PDF 文件,提取文本
        
        Args:
            file_path: PDF 文件路径
            
        Returns:
            str: 提取的文本内容
            
        Raises:
            Exception: 解析失败时抛出异常
        """
        try:
            reader = PdfReader(file_path)
            text = []
            
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text.append(page_text)
            
            return '\n\n'.join(text)
        except Exception as e:
            raise Exception(f"PDF 解析失败: {str(e)}")
    
    @staticmethod
    def parse_word(file_path):
        """
        解析 Word 文档,提取文本
        
        Args:
            file_path: Word 文件路径(.docx)
            
        Returns:
            str: 提取的文本内容
            
        Raises:
            Exception: 解析失败时抛出异常
        """
        try:
            doc = Document(file_path)
            text = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text.append(paragraph.text)
            
            # 也提取表格内容
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    text.append(' | '.join(row_text))
            
            return '\n\n'.join(text)
        except Exception as e:
            raise Exception(f"Word 解析失败: {str(e)}")
    
    @staticmethod
    def parse_file(file_path):
        """
        根据文件类型自动解析
        
        Args:
            file_path: 文件路径
            
        Returns:
            str: 提取的文本内容
        """
        ext = os.path.splitext(file_path)[1].lower()
        
        if ext == '.pdf':
            return DocumentParser.parse_pdf(file_path)
        elif ext in ['.docx', '.doc']:
            return DocumentParser.parse_word(file_path)
        else:
            raise Exception(f"不支持的文件类型: {ext}")
